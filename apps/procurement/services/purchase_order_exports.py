"""Purchase Order PDF/DOCX exports with ordered supporting attachments."""

from __future__ import annotations

import html
import re
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

from django.utils.html import strip_tags
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from openpyxl import load_workbook
from PIL import Image as PILImage
from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.utils import ImageReader

JARMO_NAME = 'Jarmo Suominen'
JARMO_TITLE = 'Sr. Vice President, Middle East\nCEO, Rejlers Abu Dhabi'
JARMO_COMPANY = 'Rejlers International Engineering Solutions AB'
COMPANY_NAME = 'Rejlers International Engineering Solutions'
COMPANY_ADDRESS = (
    'Rejlers Tower, 13th floor, AI Hamdan Street, P.O. Box 39317, '
    'Abu Dhabi, United Arab Emirates'
)
COMPANY_PHONE = '+971 50 560 6987'
COMPANY_WEBSITE = 'www.rejlers.ae'
BRAND_BLUE = colors.HexColor('#0870aa')
BRAND_TEXT_BLUE = colors.HexColor('#3275b6')
BRAND_NAVY = colors.HexColor('#1f2d55')
LOGO_PATH = Path(__file__).resolve().parent.parent / 'assets' / 'rejlers-pr-po-logo.png'
WHITE_LOGO_PATH = Path(__file__).resolve().parent.parent / 'assets' / 'rejlers-pr-po-logo-white.png'
DEFAULT_INVOICE_ADDRESS = (
    'Attn. Mr. Aneef Thadikkarantavida\n'
    'aneef.thadikkarantavida@rejlers.ae\n'
    'cc. uae.finance@rejlers.ae\n'
    'uae.procurement@rejlers.ae\n'
    'Rejlers International Engineering\n'
    'Solutions AB\n'
    'PO Box 39317\n'
    'Abu Dhabi, UAE.\n'
    'Tel: +971 2 639 7449\n'
    'Fax: +971 2 639 7448'
)
USD_TO_AED_RATE = 3.6725


def _value(value, fallback='—'):
    rendered = str(value or '').strip()
    return rendered or fallback


def _decode_html(value):
    """Decode editor content, including values that were escaped more than once."""
    decoded = str(value or '')
    for _ in range(3):
        next_value = html.unescape(decoded)
        if next_value == decoded:
            break
        decoded = next_value
    return decoded.replace('\xa0', ' ')


def _html_blocks(value):
    """Return readable text blocks from the browser rich-text editor markup."""
    source = _decode_html(value)
    if not source.strip():
        return []

    # Preserve editor paragraphs, lists, line breaks, and table rows before
    # removing unsupported markup. Django strip_tags does not decode entities.
    source = re.sub(r'(?is)<br\s*/?>', '\n', source)
    source = re.sub(r'(?is)<li\b[^>]*>', '\n• ', source)
    source = re.sub(r'(?is)</(?:p|div|h[1-6]|li|blockquote|pre|tr)>', '\n', source)
    source = re.sub(r'(?is)</(?:td|th)>', '\t', source)
    plain = _decode_html(strip_tags(source)).replace('\r\n', '\n').replace('\r', '\n')

    blocks = []
    for raw_line in plain.split('\n'):
        line = re.sub(r'[\t \f\v]+', ' ', raw_line).strip()
        if line:
            blocks.append(line)
    return blocks


def _pdf_rich_text(value, styles):
    flowables = []
    for block in _html_blocks(value):
        is_bullet = block.startswith('• ')
        content = block[2:].strip() if is_bullet else block
        style = styles['bullet'] if is_bullet else styles['body']
        flowables.extend((
            Paragraph(escape(content), style, bulletText='•' if is_bullet else None),
            Spacer(1, 1.5 * mm),
        ))
    return flowables


def _docx_rich_text(document, value):
    for block in _html_blocks(value):
        is_bullet = block.startswith('• ')
        content = block[2:].strip() if is_bullet else block
        document.add_paragraph(content, style='List Bullet' if is_bullet else None)


def _money(value, currency):
    return f'{currency or "AED"} {float(value or 0):,.2f}'


def _date_text(value):
    if not value:
        return '—'
    if hasattr(value, 'strftime'):
        rendered = value.strftime('%d %b %Y')
    try:
        from datetime import date

        rendered = date.fromisoformat(str(value)[:10]).strftime('%d %b %Y')
    except (TypeError, ValueError):
        return str(value)
    return rendered.replace(' Sep ', ' Sept ')


def _paragraph(value, style, bold=False):
    content = escape(_value(value)).replace('\n', '<br/>')
    return Paragraph(f'<b>{content}</b>' if bold else content, style)


def _buyer_reference(order):
    contacts = getattr(order, 'contact_persons', None) or {}
    references = contacts.get('buyer_references') or []
    rendered = []
    for reference in references:
        if not reference or not reference.get('name'):
            continue
        rendered.append(' · '.join(str(reference.get(key)).strip() for key in ('name', 'designation', 'email') if reference.get(key)))
    if not rendered:
        rendered.append(' · '.join(filter(None, (
            _value(getattr(order, 'buyer_reference_pm', None), 'Richa Hannah Thomas'),
            str(getattr(order, 'buyer_reference_email', '') or '').strip(),
        ))))
    return '\n'.join(rendered)


def _items(order):
    recorded = list(order.items or [])
    if not recorded:
        subtotal = max(
            0,
            float(order.total_amount or 0)
            - float(order.tax_amount or 0)
            + float(order.discount_amount or 0),
        )
        recorded = [{
            'description': order.title or order.description,
            'quantity': 1,
            'unit_price': subtotal,
            'uom': 'LOT',
        }]
    normalized = []
    for index, item in enumerate(recorded):
        quantity = float(item.get('quantity', item.get('qty', 1)) or 0)
        unit_price = float(item.get('unit_price', item.get('price', 0)) or 0)
        discount = float(item.get('discount', 0) or 0)
        total = float(item.get('total', item.get('line_total', 0)) or 0)
        if not total:
            total = max(0, (quantity * unit_price) - discount)
        normalized.append({
            'line': item.get('line_code') or item.get('item_code') or index + 1,
            'description': item.get('description') or item.get('item') or item.get('name') or order.title,
            'specification': item.get('specification') or '',
            'comment': item.get('comment') or item.get('comments') or item.get('remarks') or item.get('notes') or '',
            'quantity': quantity,
            'uom': item.get('uom') or item.get('unit') or 'EA',
            'unit_price': unit_price,
            'discount': discount,
            'total': total,
        })
    return normalized


def _attachment(entry, index):
    if isinstance(entry, str):
        return {
            'title': f'Attachment {index + 1}',
            'description': Path(entry).name,
            'filename': Path(entry).name,
            's3_url': entry,
        }
    entry = dict(entry or {})
    return {
        **entry,
        'title': _value(entry.get('title'), f'Attachment {index + 1}'),
        'description': _value(
            entry.get('description'),
            entry.get('filename') or entry.get('name') or f'Attachment {index + 1}',
        ),
        'filename': _value(entry.get('filename') or entry.get('name'), f'attachment-{index + 1}'),
    }


def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        'title': ParagraphStyle('POTitle', parent=styles['Title'], fontSize=17, leading=21, textColor=colors.HexColor('#16689b')),
        'heading': ParagraphStyle('POHeading', parent=styles['Heading2'], fontSize=10, leading=13, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor('#1f2937')),
        'body': ParagraphStyle('POBody', parent=styles['BodyText'], fontSize=8.5, leading=11),
        'bullet': ParagraphStyle('POBullet', parent=styles['BodyText'], fontSize=8.5, leading=11, leftIndent=5 * mm, firstLineIndent=-3 * mm),
        'small': ParagraphStyle('POSmall', parent=styles['BodyText'], fontSize=7, leading=9),
        'right': ParagraphStyle('PORight', parent=styles['BodyText'], fontSize=8.5, leading=11, alignment=TA_RIGHT),
        'cover': ParagraphStyle('POCover', parent=styles['Title'], fontSize=22, leading=28, alignment=TA_CENTER, textColor=colors.HexColor('#16689b')),
        'cover_body': ParagraphStyle('POCoverBody', parent=styles['BodyText'], fontSize=12, leading=18, alignment=TA_CENTER),
        'preview': ParagraphStyle('POPreview', parent=styles['BodyText'], fontSize=6.8, leading=8.3, textColor=colors.HexColor('#334155')),
        'preview_bold': ParagraphStyle('POPreviewBold', parent=styles['BodyText'], fontSize=6.8, leading=8.3, fontName='Helvetica-Bold', textColor=colors.HexColor('#334155')),
        'preview_heading': ParagraphStyle('POPreviewHeading', parent=styles['Heading2'], fontSize=7.8, leading=10, fontName='Helvetica-Bold', textColor=colors.HexColor('#1f2937')),
    }


def _draw_rejlers_wordmark(canvas, x, y, width, color):
    """Draw the official wordmark, with a vector fallback for damaged installs."""
    logo_path = WHITE_LOGO_PATH if color == colors.white else LOGO_PATH
    if logo_path.exists():
        canvas.drawImage(
            ImageReader(str(logo_path)), x, y, width=width, height=width / 6.64,
            preserveAspectRatio=True, mask='auto', anchor='sw',
        )
        return
    scale = width / 42.0
    canvas.saveState()
    canvas.setStrokeColor(color)
    canvas.setLineWidth(max(0.65, 1.15 * scale))
    canvas.setLineCap(1)
    canvas.setLineJoin(1)
    canvas.line(x, y + 1.0 * scale, x + 4.2 * scale, y + 5.2 * scale)
    canvas.line(x + 4.2 * scale, y + 5.2 * scale, x + 4.2 * scale, y + 1.0 * scale)
    canvas.setFillColor(color)
    canvas.setFont('Helvetica', max(4.5, 5.8 * scale))
    canvas.drawString(x + 6.2 * scale, y, 'REJLERS')
    canvas.restoreState()


def _pdf_page(canvas, document, order, page_number=None):
    """Draw the same branded header/footer used by Print Preview."""
    canvas.saveState()
    width, height = A4

    # Header: document identity on the left, Rejlers wordmark on the right.
    left = 16 * mm
    top = height - 11 * mm
    canvas.setFillColor(BRAND_TEXT_BLUE)
    canvas.setFont('Helvetica-Bold', 8.5)
    canvas.drawString(left, top, 'PURCHASE ORDER')
    canvas.setFont('Helvetica-Bold', 7.5)
    canvas.drawString(left, top - 4 * mm, _value(order.po_number, 'PO NUMBER PENDING'))
    canvas.setFillColor(colors.HexColor('#64748b'))
    canvas.setFont('Helvetica', 5.5)
    canvas.drawString(left, top - 7 * mm, _value(getattr(order, 'form_note', None), '(PO no. to be used in all documents)'))
    canvas.setFillColor(BRAND_TEXT_BLUE)
    canvas.setFont('Helvetica-Bold', 7)
    canvas.drawString(left, top - 12 * mm, _date_text(getattr(order, 'po_date', None)))

    logo_width = 27 * mm
    logo_x = width - left - logo_width
    _draw_rejlers_wordmark(canvas, logo_x, top - 1.5 * mm, logo_width, BRAND_NAVY)
    canvas.setFillColor(BRAND_TEXT_BLUE)
    canvas.setFont('Helvetica-Bold', 5.7)
    canvas.drawRightString(width - left, top - 9 * mm, 'HOME OF THE')
    canvas.drawRightString(width - left, top - 12 * mm, 'LEARNING MINDS')

    # Footer: repeated white brand marks in the blue band, then the same
    # company/contact block and page number shown by the browser preview.
    band_x = left
    band_y = 12 * mm
    band_width = width - (2 * left)
    band_height = 7 * mm
    canvas.setFillColor(BRAND_BLUE)
    canvas.rect(band_x, band_y, band_width, band_height, fill=1, stroke=0)
    group_width = band_width / 5
    for index in (0, 2, 4):
        footer_logo_width = 21 * mm
        _draw_rejlers_wordmark(
            canvas,
            band_x + ((index + 0.5) * group_width) - (footer_logo_width / 2),
            band_y + 1.9 * mm,
            footer_logo_width,
            colors.white,
        )
    canvas.setFillColor(colors.white)
    canvas.setFont('Helvetica-Bold', 4.5)
    for index in (1, 3):
        center = band_x + ((index + 0.5) * group_width)
        canvas.drawCentredString(center, band_y + 4.2 * mm, 'HOME of the')
        canvas.drawCentredString(center, band_y + 2.2 * mm, 'LEARNING MINDS')

    canvas.setFillColor(BRAND_TEXT_BLUE)
    canvas.setFont('Helvetica', 4.6)
    canvas.drawString(left + 8 * mm, 9.2 * mm, COMPANY_NAME)
    canvas.drawString(left + 8 * mm, 7.0 * mm, COMPANY_ADDRESS)
    canvas.drawString(left + 8 * mm, 4.8 * mm, f'Tel: {COMPANY_PHONE} | {COMPANY_WEBSITE}')
    canvas.setFont('Helvetica', 5.2)
    canvas.drawRightString(width - left, 5.5 * mm, f'Page {page_number or document.page}')
    canvas.restoreState()


def _main_pdf(order):
    output = BytesIO()
    styles = _pdf_styles()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=34 * mm,
        bottomMargin=25 * mm,
        title=_value(order.po_number),
    )
    currency = order.currency or 'AED'
    items = _items(order)
    subtotal = sum(item['total'] for item in items)
    tax = float(order.tax_amount or 0)
    total = float(order.total_amount or subtotal + tax)
    vendor = getattr(order, 'vendor', None)
    preview = styles['preview']
    preview_bold = styles['preview_bold']

    def pair_rows(rows):
        return Table(
            [[Paragraph(f'<b>{escape(label)}:</b>', preview), _paragraph(value, preview, strong)] for label, value, strong in rows],
            colWidths=[24 * mm, 58 * mm],
            style=TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 1.5 * mm),
                ('TOPPADDING', (0, 0), (-1, -1), 1.2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1.8),
            ]),
        )

    invoice_address = DEFAULT_INVOICE_ADDRESS
    if getattr(order, 'invoicing_attn', None) or getattr(order, 'invoicing_emails', None):
        invoice_emails = getattr(order, 'invoicing_emails', None) or []
        if not isinstance(invoice_emails, (list, tuple)):
            invoice_emails = [invoice_emails]
        invoice_address = '\n'.join(filter(None, (
            str(getattr(order, 'invoicing_attn', '') or '').strip(),
            *[str(email).strip() for email in invoice_emails if email],
            'Rejlers International Engineering Solutions AB',
            'PO Box 39317', 'Abu Dhabi, UAE.', 'Tel: +971 2 639 7449',
            f'Fax: {_value(getattr(order, "company_fax", None), "+971 2 639 7448")}',
        )))

    details = Table([[pair_rows([
        ('Seller information', '\n'.join(filter(None, (
            _value(getattr(vendor, 'name', None)),
            str(getattr(order, 'seller_address', '') or '').strip(),
        ))), False),
        ('Invoicing Address', invoice_address, False),
    ]), '', pair_rows([
        ('Seller Reference', getattr(order, 'seller_reference', None), False),
        ('Quote Ref.', getattr(order, 'quote_ref', None), False),
        ('License No.', getattr(order, 'seller_license_no', None), False),
        ('Buyer Reference', _buyer_reference(order), False),
    ])]], colWidths=[84 * mm, 8 * mm, 84 * mm], style=TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    commercial = Table([[pair_rows([
        ('Payment Terms', getattr(order, 'payment_terms', None), False),
        ('Payment Mode', getattr(order, 'payment_mode', None), False),
        ('Project', getattr(order, 'project_number', None) or getattr(order, 'rad_project_no', None) or 'Multiple Projects', True),
    ]), '', pair_rows([
        ('Delivery terms', getattr(order, 'delivery_terms', None), False),
        ('Delivery date', _date_text(getattr(order, 'expected_delivery', None)), False),
        ('Marking', getattr(order, 'marking', None) or order.po_number, True),
    ])]], colWidths=[84 * mm, 8 * mm, 84 * mm], style=TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    summary_table = Table([[
        Paragraph(f'<b>Purchase Summary:</b><br/><br/><b>{escape(_value(getattr(order, "summary", None) or order.title))}</b>', preview),
        '',
        Table([
            [Paragraph('<b>Total Purchase Price:</b>', preview), Paragraph(escape(f'{subtotal:,.2f} {currency}'), styles['right'])],
            [Paragraph(f'<b>VAT ({float(getattr(order, "vat_percentage", 0) or 0):g}%):</b>', preview), Paragraph(escape(f'{tax:,.2f} {currency}'), styles['right'])],
            [Paragraph('<b>Total Sum:</b>', preview_bold), Paragraph(escape(f'{total:,.2f} {currency}'), styles['right'])],
        ], colWidths=[45 * mm, 34 * mm], style=TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 1), ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
        ])),
    ]], colWidths=[91 * mm, 6 * mm, 79 * mm], style=TableStyle([
        ('LINEABOVE', (0, 0), (-1, 0), 1.2, colors.HexColor('#475569')),
        ('LINEBELOW', (0, 0), (-1, 0), 1.2, colors.HexColor('#475569')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 2.2 * mm), ('BOTTOMPADDING', (0, 0), (-1, -1), 2.2 * mm),
    ]))
    approval_name = _value(getattr(order, 'approved_by_name', None), JARMO_NAME)
    approval_title = _value(getattr(order, 'approved_by_title', None), JARMO_TITLE)
    approved = Paragraph(
        '<b>Approved by:</b><br/><br/><br/><br/><br/><br/>'
        f'<b>{escape(approval_name)}</b><br/>{escape(approval_title).replace(chr(10), "<br/>")}<br/>'
        f'{JARMO_COMPANY}<br/><b>Date:</b> {escape(_value(getattr(order, "approved_date", None), ""))}', preview,
    )
    confirmation = Paragraph(
        '<b>Order Confirmation:</b><br/>We acknowledge receipt of your documents and will perform according to this PO.'
        '<br/><br/><b>Seller Signature:</b> ______________________________'
        f'<br/><br/><b>Date:</b> {_date_text(getattr(order, "confirmation_date", None))}'
        f'<br/><br/><b>Seller information:</b> {escape(_value(getattr(vendor, "name", None)))}'
        f'<br/><br/><b>Phone / Email:</b> {escape(" / ".join(filter(None, (str(getattr(order, "seller_phone", "") or ""), str(getattr(order, "seller_email", "") or "")))) or "â€”")}',
        preview,
    )
    approval_table = Table([[approved, '', confirmation]], colWidths=[86 * mm, 5 * mm, 85 * mm], style=TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LINEBEFORE', (2, 0), (2, 0), 0.5, colors.HexColor('#64748b')),
        ('LEFTPADDING', (0, 0), (0, 0), 0), ('RIGHTPADDING', (0, 0), (0, 0), 7 * mm),
        ('LEFTPADDING', (2, 0), (2, 0), 3 * mm), ('RIGHTPADDING', (2, 0), (2, 0), 0),
    ]))
    story = [
        Spacer(1, 2 * mm), details, Spacer(1, 5 * mm), commercial, Spacer(1, 4 * mm),
        summary_table, Spacer(1, 8 * mm), approval_table, PageBreak(),
        Paragraph(f'<u>PURCHASE ORDER:</u> &nbsp;{escape(_value(order.title))}', styles['heading']),
        Paragraph(
            f'We, {COMPANY_NAME} (Buyer), issue this purchase order to '
            f'<b>{escape(_value(getattr(vendor, "name", None)))}</b> (Seller).', preview,
        ),
        Paragraph('PO DESCRIPTION &amp; SCOPE', styles['heading']),
    ]
    narrative = _pdf_rich_text(order.description, styles)
    story.extend(narrative or [Paragraph(escape(_value(order.title)), styles['body'])])
    # Match the live A4 document: the price summary starts on a clean page.
    # This also prevents an orphaned heading or split table after long scope text.
    story.extend((PageBreak(), Paragraph('SUMMARY OF PRICES', styles['heading'])))
    price_rows = [[
        Paragraph('<b>Line Code</b>', styles['small']),
        Paragraph('<b>Item Description</b>', styles['small']),
        Paragraph('<b>Comment</b>', styles['small']),
        Paragraph('<b>Qty.</b>', styles['small']),
        Paragraph('<b>UOM</b>', styles['small']),
        Paragraph('<b>Unit Price</b>', styles['small']),
        Paragraph('<b>Discount</b>', styles['small']),
        Paragraph('<b>Total Price</b>', styles['small']),
    ]]
    for item in items:
        price_rows.append([
            Paragraph(escape(str(item['line'])), styles['small']),
            Paragraph(escape(_value(item['description'])), styles['small']),
            Paragraph(escape(_value(item['comment'], '')), styles['small']),
            Paragraph(f'{item["quantity"]:g}', styles['right']),
            Paragraph(escape(_value(item['uom'], '')), styles['small']),
            Paragraph(escape(_money(item['unit_price'], currency)), styles['right']),
            Paragraph(escape(_money(item['discount'], currency)), styles['right']),
            Paragraph(escape(_money(item['total'], currency)), styles['right']),
        ])
    story.extend([
        Table(price_rows, repeatRows=1, colWidths=[11 * mm, 55 * mm, 32 * mm, 9 * mm, 10 * mm, 20 * mm, 16 * mm, 23 * mm], style=TableStyle([
            ('LINEABOVE', (0, 0), (-1, 0), 1.2, colors.HexColor('#475569')),
            ('LINEBELOW', (0, 0), (-1, 0), 1.2, colors.HexColor('#475569')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ])),
        Spacer(1, 6 * mm),
        Table([
            ['Total Price:', _money(subtotal, currency)],
            [f'VAT ({float(order.vat_percentage or 0):g}%):', _money(tax, currency)],
            ['Total Sum:', _money(total, currency)],
            *([['Grand Total USD in AED:', _money(total * USD_TO_AED_RATE, 'AED')]] if str(currency).upper() == 'USD' else []),
        ], colWidths=[55 * mm, 42 * mm], hAlign='RIGHT', style=TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#475569')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTSIZE', (0, 0), (-1, -1), 8.5),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#e2e8f0')),
        ])),
    ])
    document.build(
        story,
        onFirstPage=lambda canvas, doc: _pdf_page(canvas, doc, order),
        onLaterPages=lambda canvas, doc: _pdf_page(canvas, doc, order),
    )
    return output.getvalue()


def _cover_pdf(order, attachment, index, page_number):
    output = BytesIO()
    styles = _pdf_styles()
    document = SimpleDocTemplate(output, pagesize=A4, leftMargin=25 * mm, rightMargin=25 * mm, topMargin=35 * mm, bottomMargin=25 * mm)
    document.build([
        Spacer(1, 55 * mm),
        Paragraph(f'ATTACHMENT - {index + 1}', styles['cover']),
        Spacer(1, 12 * mm),
        Paragraph(escape(_value(attachment['description'])), styles['cover_body']),
    ], onFirstPage=lambda canvas, doc: _pdf_page(canvas, doc, order, page_number))
    return output.getvalue()


def _download_attachment(attachment):
    from django.conf import settings
    from django.core.files.storage import default_storage

    key = str(attachment.get('s3_key') or '').strip()
    if not key:
        return None
    try:
        with default_storage.open(key, 'rb') as stored_file:
            return stored_file.read()
    except Exception:
        # Older PO uploads used a raw bucket key without the MediaStorage
        # prefix. Keep those files exportable after adopting default_storage.
        if not getattr(settings, 'USE_S3', False):
            return None
        try:
            from apps.core.s3_utils import S3Client

            output = BytesIO()
            if S3Client().download_file(key, output):
                return output.getvalue()
        except Exception:
            return None
    return None


def _flowables_pdf(flowables):
    output = BytesIO()
    SimpleDocTemplate(output, pagesize=A4, leftMargin=14 * mm, rightMargin=14 * mm, topMargin=14 * mm, bottomMargin=14 * mm).build(flowables)
    return output.getvalue()


def _image_pdf(content):
    source = BytesIO(content)
    with PILImage.open(source) as image:
        width, height = image.size
    max_width, max_height = 180 * mm, 267 * mm
    scale = min(max_width / width, max_height / height)
    source.seek(0)
    rendered = Image(source, width=width * scale, height=height * scale)
    rendered.hAlign = 'CENTER'
    return _flowables_pdf([rendered])


def _docx_pdf(content):
    document = Document(BytesIO(content))
    styles = _pdf_styles()
    flowables = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            flowables.extend([Paragraph(escape(value), styles['body']), Spacer(1, 2 * mm)])
    for table in document.tables:
        rows = [[Paragraph(escape(cell.text), styles['small']) for cell in row.cells] for row in table.rows]
        if rows:
            flowables.extend([Table(rows, repeatRows=1, style=TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.35, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ])), Spacer(1, 3 * mm)])
    return _flowables_pdf(flowables or [Paragraph('The attached Word document has no renderable text.', styles['body'])])


def _xlsx_pdf(content):
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    styles = _pdf_styles()
    flowables = []
    for sheet in workbook.worksheets:
        flowables.append(Paragraph(escape(sheet.title), styles['heading']))
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = [Paragraph(escape(_value(cell, '')), styles['small']) for cell in row]
            if any(str(cell or '').strip() for cell in row):
                rows.append(values)
            if len(rows) >= 100:
                break
        if rows:
            available = 180 * mm
            columns = max(len(row) for row in rows)
            rows = [row + [Paragraph('', styles['small'])] * (columns - len(row)) for row in rows]
            flowables.extend([Table(rows, repeatRows=1, colWidths=[available / columns] * columns, style=TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ])), PageBreak()])
    if flowables and isinstance(flowables[-1], PageBreak):
        flowables.pop()
    return _flowables_pdf(flowables or [Paragraph('The attached workbook has no renderable cells.', styles['body'])])


def _attachment_pdf(content, filename, content_type=''):
    suffix = Path(filename).suffix.lower()
    if suffix == '.pdf' or content_type == 'application/pdf':
        PdfReader(BytesIO(content))
        return content
    if suffix in {'.png', '.jpg', '.jpeg'} or str(content_type).startswith('image/'):
        return _image_pdf(content)
    if suffix == '.docx':
        return _docx_pdf(content)
    if suffix == '.xlsx':
        return _xlsx_pdf(content)
    return None


def build_purchase_order_pdf(order):
    """Return one PDF containing the PO, attachment covers, and renderable files."""
    writer = PdfWriter()
    warnings = []

    def append(pdf_bytes):
        for page in PdfReader(BytesIO(pdf_bytes)).pages:
            writer.add_page(page)

    append(_main_pdf(order))
    for index, raw_attachment in enumerate(order.attachments or []):
        attachment = _attachment(raw_attachment, index)
        append(_cover_pdf(order, attachment, index, len(writer.pages) + 1))
        content = _download_attachment(attachment)
        if content is None:
            warnings.append(f'{attachment["filename"]}: file could not be downloaded')
            continue
        try:
            rendered = _attachment_pdf(content, attachment['filename'], attachment.get('content_type'))
            if rendered:
                append(rendered)
            else:
                warnings.append(f'{attachment["filename"]}: format cannot be rendered in PDF')
        except Exception as exc:  # The PO must remain downloadable if one supporting file is corrupt.
            warnings.append(f'{attachment["filename"]}: {type(exc).__name__}')

    output = BytesIO()
    writer.write(output)
    return output.getvalue(), warnings


def _docx_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        properties.append(shading)
    shading.set(qn('w:fill'), fill)


def _docx_set_cell_text(cell, value, *, size=7, bold=False, color='334155', align=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(str(value or ''))
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return paragraph


def _docx_add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' PAGE '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend((begin, instruction, end))


def _docx_no_borders(table):
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        properties.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'nil')
        borders.append(element)


def _configure_docx_header_footer(document, order):
    section = document.sections[0]
    section.header_distance = Mm(7)
    section.footer_distance = Mm(5)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Mm(178))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(135)
    table.columns[1].width = Mm(43)
    left, right = table.rows[0].cells
    paragraph = _docx_set_cell_text(left, 'PURCHASE ORDER', size=8, bold=True, color='3275B6')
    for value, size, bold, color in (
        (_value(order.po_number, 'PO NUMBER PENDING'), 7, True, '3275B6'),
        (_value(getattr(order, 'form_note', None), '(PO no. to be used in all documents)'), 5.5, False, '64748B'),
        (_date_text(getattr(order, 'po_date', None)), 7, True, '3275B6'),
    ):
        run = paragraph.add_run(f'\n{value}')
        run.bold = bold
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
    right.text = ''
    logo_paragraph = right.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_paragraph.paragraph_format.space_after = Pt(0)
    if LOGO_PATH.exists():
        logo_paragraph.add_run().add_picture(str(LOGO_PATH), width=Mm(27))
    tagline = logo_paragraph.add_run('\nHOME OF THE\nLEARNING MINDS')
    tagline.bold = True
    tagline.font.name = 'Arial'
    tagline.font.size = Pt(6.5)
    tagline.font.color.rgb = RGBColor.from_string('3275B6')

    footer = section.footer
    band = footer.add_table(rows=1, cols=5, width=Mm(178))
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = ('REJLERS', 'HOME of the\nLEARNING MINDS', 'REJLERS', 'HOME of the\nLEARNING MINDS', 'REJLERS')
    for index, (cell, value) in enumerate(zip(band.rows[0].cells, values)):
        _docx_cell_shading(cell, '0870AA')
        paragraph = _docx_set_cell_text(cell, value, size=5.5, bold=True, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)
        if index in (0, 2, 4) and WHITE_LOGO_PATH.exists():
            paragraph.clear()
            paragraph.add_run().add_picture(str(WHITE_LOGO_PATH), height=Mm(3.2))
    details = footer.add_table(rows=1, cols=2, width=Mm(160))
    details.alignment = WD_TABLE_ALIGNMENT.CENTER
    _docx_set_cell_text(
        details.cell(0, 0),
        f'{COMPANY_NAME}\n{COMPANY_ADDRESS}\nTel: {COMPANY_PHONE} | {COMPANY_WEBSITE}',
        size=4.5, color='4E83AD',
    )
    page_paragraph = _docx_set_cell_text(details.cell(0, 1), 'Page ', size=5, color='4E83AD', align=WD_ALIGN_PARAGRAPH.RIGHT)
    _docx_add_page_field(page_paragraph)


def build_purchase_order_docx(order):
    """Return the editable PO through Summary of Prices; attachments are excluded."""
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(29)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    _configure_docx_header_footer(document, order)

    vendor = getattr(order, 'vendor', None)
    invoice_emails = getattr(order, 'invoicing_emails', None) or []
    if not isinstance(invoice_emails, (list, tuple)):
        invoice_emails = [invoice_emails]
    invoice_address = '\n'.join(filter(None, (
        str(getattr(order, 'invoicing_attn', '') or '').strip(),
        *[str(email).strip() for email in invoice_emails if email],
    ))) or DEFAULT_INVOICE_ADDRESS
    details = document.add_table(rows=4, cols=4)
    _docx_no_borders(details)
    detail_rows = (
        ('Seller information', '\n'.join(filter(None, (_value(getattr(vendor, 'name', None)), str(getattr(order, 'seller_address', '') or '').strip()))), 'Seller Reference', getattr(order, 'seller_reference', None)),
        ('Invoicing Address', invoice_address, 'Quote Ref.', getattr(order, 'quote_ref', None)),
        ('', '', 'License No.', getattr(order, 'seller_license_no', None)),
        ('', '', 'Buyer Reference', _buyer_reference(order)),
    )
    for row, values in zip(details.rows, detail_rows):
        for index, value in enumerate(values):
            _docx_set_cell_text(row.cells[index], _value(value, '') if index % 2 else value, bold=index % 2 == 0)

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    commercial = document.add_table(rows=3, cols=4)
    _docx_no_borders(commercial)
    commercial_rows = (
        ('Payment Terms', getattr(order, 'payment_terms', None), 'Delivery terms', getattr(order, 'delivery_terms', None)),
        ('Payment Mode', getattr(order, 'payment_mode', None), 'Delivery date', _date_text(getattr(order, 'expected_delivery', None))),
        ('Project', getattr(order, 'project_number', None) or getattr(order, 'rad_project_no', None) or 'Multiple Projects', 'Marking', getattr(order, 'marking', None) or order.po_number),
    )
    for row, values in zip(commercial.rows, commercial_rows):
        for index, value in enumerate(values):
            _docx_set_cell_text(row.cells[index], _value(value), bold=index % 2 == 0 or (index == 1 and row is commercial.rows[-1]))

    summary = document.add_table(rows=1, cols=2)
    summary.style = 'Table Grid'
    _docx_set_cell_text(summary.cell(0, 0), f'Purchase Summary:\n{_value(getattr(order, "summary", None) or order.title)}', size=7, bold=True)
    subtotal_for_summary = sum(item['total'] for item in _items(order))
    _docx_set_cell_text(
        summary.cell(0, 1),
        f'Total Purchase Price: {subtotal_for_summary:,.2f} {order.currency or "AED"}\n'
        f'VAT ({float(order.vat_percentage or 0):g}%): {float(order.tax_amount or 0):,.2f} {order.currency or "AED"}\n'
        f'Total Sum: {float(order.total_amount or 0):,.2f} {order.currency or "AED"}',
        size=7, bold=True,
    )

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    approval = document.add_table(rows=1, cols=2)
    _docx_no_borders(approval)
    _docx_set_cell_text(
        approval.cell(0, 0),
        f'Approved by:\n\n\n\n\n{_value(getattr(order, "approved_by_name", None), JARMO_NAME)}\n'
        f'{_value(getattr(order, "approved_by_title", None), JARMO_TITLE)}\n{JARMO_COMPANY}\n'
        f'Date: {_value(getattr(order, "approved_date", None), "")}', size=7,
    )
    _docx_set_cell_text(
        approval.cell(0, 1),
        f'Order Confirmation:\nWe acknowledge receipt of your documents and will perform according to this PO.\n\n'
        f'Seller Signature: ____________________\n\nDate: {_date_text(getattr(order, "confirmation_date", None))}\n\n'
        f'Seller information: {_value(getattr(vendor, "name", None))}\n\n'
        f'Phone / Email: {" / ".join(filter(None, (str(getattr(order, "seller_phone", "") or ""), str(getattr(order, "seller_email", "") or "")))) or "â€”"}', size=7,
    )
    document.add_page_break()
    document.add_heading(f'PURCHASE ORDER:  {_value(order.title)}', level=1)
    document.add_paragraph(f'We, {COMPANY_NAME} (Buyer), issue this purchase order to {_value(getattr(vendor, "name", None))} (Seller).')
    document.add_heading('PO DESCRIPTION & SCOPE', level=1)
    if _html_blocks(order.description):
        _docx_rich_text(document, order.description)
    else:
        document.add_paragraph(_value(order.title))
    document.add_page_break()
    document.add_heading('Summary of Prices', level=1)
    items = _items(order)
    table = document.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    for cell, heading in zip(table.rows[0].cells, ('Line Code', 'Item Description', 'Comment', 'Qty.', 'UOM', 'Unit Price', 'Discount', 'Total Price')):
        cell.text = heading
    currency = order.currency or 'AED'
    for item in items:
        cells = table.add_row().cells
        values = (
            item['line'], item['description'], item['comment'], f'{item["quantity"]:g}', item['uom'],
            _money(item['unit_price'], currency), _money(item['discount'], currency), _money(item['total'], currency),
        )
        for cell, value in zip(cells, values):
            cell.text = str(value)
    subtotal = sum(item['total'] for item in items)
    tax = float(order.tax_amount or 0)
    total = float(order.total_amount or subtotal + tax)
    totals = document.add_paragraph()
    totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_text = (
        f'Total Price: {_money(subtotal, currency)}\n'
        f'VAT ({float(order.vat_percentage or 0):g}%): {_money(tax, currency)}\n'
        f'Total Sum: {_money(total, currency)}'
    )
    if str(currency).upper() == 'USD':
        totals_text += f'\nGrand Total USD in AED: {_money(total * USD_TO_AED_RATE, "AED")}'
    totals.add_run(totals_text).bold = True
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(9)
    output = BytesIO()
    document.save(output)
    return output.getvalue()

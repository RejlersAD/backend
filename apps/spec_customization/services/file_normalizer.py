"""
File Normalizer — convert any uploaded Paper Spec input into a PDF.

Smart, soft-coded multi-format ingest: PDFs pass through, images are wrapped
into a single-page PDF, Office files (DOCX/XLSX/XLSM) are rendered to PDF via
text-extraction + ReportLab, and plain text is rendered to a paginated PDF.

All format knobs live in `SUPPORTED_FORMATS` below — add a new key to teach
the platform a new format.

Returns a `NormalizedFile` with:
    - django File (in-memory ContentFile)
    - new filename (PDF)
    - mime  ('application/pdf')
    - original_extension, original_mime, conversion_engine, conversion_notes

Never raises for unsupported formats — instead returns `success=False` and a
human-readable message that the view can surface to the user.
"""
from __future__ import annotations

import io
import logging
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)

# ─── Soft-coded format registry ──────────────────────────────────────────────
SUPPORTED_FORMATS = {
    # PDFs — passthrough
    'pdf':  {'mime': 'application/pdf',                                'engine': 'passthrough',  'group': 'pdf'},
    # Raster images — wrap to single-page PDF
    'png':  {'mime': 'image/png',                                      'engine': 'image_to_pdf', 'group': 'image'},
    'jpg':  {'mime': 'image/jpeg',                                     'engine': 'image_to_pdf', 'group': 'image'},
    'jpeg': {'mime': 'image/jpeg',                                     'engine': 'image_to_pdf', 'group': 'image'},
    'tif':  {'mime': 'image/tiff',                                     'engine': 'image_to_pdf', 'group': 'image'},
    'tiff': {'mime': 'image/tiff',                                     'engine': 'image_to_pdf', 'group': 'image'},
    'bmp':  {'mime': 'image/bmp',                                      'engine': 'image_to_pdf', 'group': 'image'},
    'webp': {'mime': 'image/webp',                                     'engine': 'image_to_pdf', 'group': 'image'},
    'gif':  {'mime': 'image/gif',                                      'engine': 'image_to_pdf', 'group': 'image'},
    # Word
    'docx': {'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
             'engine': 'docx_to_pdf', 'group': 'office'},
    'doc':  {'mime': 'application/msword',                             'engine': 'docx_to_pdf', 'group': 'office'},
    # Excel
    'xlsx': {'mime': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
             'engine': 'xlsx_to_pdf', 'group': 'office'},
    'xlsm': {'mime': 'application/vnd.ms-excel.sheet.macroEnabled.12', 'engine': 'xlsx_to_pdf', 'group': 'office'},
    'xls':  {'mime': 'application/vnd.ms-excel',                       'engine': 'xlsx_to_pdf', 'group': 'office'},
    # Plain text / CSV
    'txt':  {'mime': 'text/plain',                                     'engine': 'text_to_pdf', 'group': 'text'},
    'csv':  {'mime': 'text/csv',                                       'engine': 'text_to_pdf', 'group': 'text'},
    'log':  {'mime': 'text/plain',                                     'engine': 'text_to_pdf', 'group': 'text'},
}

# A few additional knobs (kept here so they are easy to tune).
NORMALIZER_CONFIG = {
    'image_page_size_pt':       (595, 842),   # A4 portrait
    'image_dpi_for_pdf':        150,
    'text_page_size_pt':        (595, 842),
    'text_margin_pt':           40,
    'text_font_name':           'Helvetica',
    'text_font_size':           9,
    'text_line_height':         11,
    'text_max_chars_per_line':  100,
    'xlsx_max_rows_per_sheet':  500,         # safety cap to avoid 100k-row PDFs
    'xlsx_max_cols':            20,
    'xlsx_col_width_chars':     14,
}


@dataclass
class NormalizedFile:
    success: bool
    file: Optional[ContentFile] = None
    filename: str = ''
    mime: str = ''
    original_extension: str = ''
    original_mime: str = ''
    conversion_engine: str = ''
    conversion_notes: str = ''
    error_message: str = ''
    warnings: List[str] = field(default_factory=list)


def _ext(name: str) -> str:
    return (os.path.splitext(name or '')[1] or '').lower().lstrip('.')


def get_accepted_extensions() -> List[str]:
    """Sorted list of all supported input extensions (for frontend mirror)."""
    return sorted(SUPPORTED_FORMATS.keys())


def detect_format(filename: str) -> dict:
    """Return descriptor for the given filename, or empty dict."""
    return SUPPORTED_FORMATS.get(_ext(filename), {})


# ─── Conversion engines ──────────────────────────────────────────────────────
def _image_to_pdf(raw: bytes, src_ext: str) -> bytes:
    """Wrap one image into a single-page PDF that fits A4."""
    from PIL import Image
    page_w, page_h = NORMALIZER_CONFIG['image_page_size_pt']

    img = Image.open(io.BytesIO(raw))
    # Multi-frame TIFF — flatten to first frame for simplicity (Pillow supports save_all)
    if getattr(img, 'is_animated', False) or src_ext in ('tif', 'tiff'):
        frames = []
        try:
            i = 0
            while True:
                img.seek(i)
                frames.append(img.convert('RGB').copy())
                i += 1
        except EOFError:
            pass
        if not frames:
            frames = [img.convert('RGB')]
    else:
        frames = [img.convert('RGB')]

    out = io.BytesIO()
    first, rest = frames[0], frames[1:]
    first.save(out, format='PDF', save_all=True, append_images=rest, resolution=NORMALIZER_CONFIG['image_dpi_for_pdf'])
    return out.getvalue()


def _text_to_pdf(text: str, title: str = '') -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    page_w, page_h = A4
    margin       = NORMALIZER_CONFIG['text_margin_pt']
    line_h       = NORMALIZER_CONFIG['text_line_height']
    font_name    = NORMALIZER_CONFIG['text_font_name']
    font_size    = NORMALIZER_CONFIG['text_font_size']
    max_chars    = NORMALIZER_CONFIG['text_max_chars_per_line']

    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=A4)
    if title:
        c.setTitle(title)

    def wrap(line: str) -> List[str]:
        if not line:
            return ['']
        return [line[i:i + max_chars] for i in range(0, len(line), max_chars)]

    y = page_h - margin
    c.setFont(font_name, font_size)
    for raw_line in (text or '').splitlines() or ['']:
        for piece in wrap(raw_line):
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = page_h - margin
            c.drawString(margin, y, piece)
            y -= line_h
    c.save()
    return out.getvalue()


def _docx_to_pdf(raw: bytes) -> bytes:
    """Extract text from DOCX (paragraphs + tables) and render as PDF."""
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(raw))
    lines: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            lines.append(p.text)
    for tbl in doc.tables:
        lines.append('')
        for row in tbl.rows:
            cells = [(cell.text or '').replace('\n', ' ').strip() for cell in row.cells]
            lines.append(' | '.join(cells))
        lines.append('')
    return _text_to_pdf('\n'.join(lines))


def _xlsx_to_pdf(raw: bytes) -> bytes:
    """Dump each sheet as pipe-separated rows (capped) and render as PDF."""
    from openpyxl import load_workbook
    max_rows = NORMALIZER_CONFIG['xlsx_max_rows_per_sheet']
    max_cols = NORMALIZER_CONFIG['xlsx_max_cols']

    wb = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    lines: List[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f'=== Sheet: {sheet_name} ===')
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                lines.append(f'… (truncated at {max_rows} rows)')
                break
            cells = ['' if v is None else str(v) for v in row[:max_cols]]
            lines.append(' | '.join(cells))
        lines.append('')
    return _text_to_pdf('\n'.join(lines))


# ─── Public entry point ──────────────────────────────────────────────────────
def normalize_to_pdf(django_file) -> NormalizedFile:
    """
    Convert a Django UploadedFile of any supported format into a PDF.

    Always returns a NormalizedFile. Caller should check `.success` and
    surface `.error_message` to the user if False.
    """
    if not django_file:
        return NormalizedFile(success=False, error_message='No file provided.')

    src_name = getattr(django_file, 'name', '') or 'upload'
    src_ext  = _ext(src_name)
    src_mime = getattr(django_file, 'content_type', '') or ''

    desc = SUPPORTED_FORMATS.get(src_ext)
    if not desc:
        accepted = ', '.join(get_accepted_extensions())
        return NormalizedFile(
            success=False,
            original_extension=src_ext,
            original_mime=src_mime,
            error_message=(
                f"Unsupported format '.{src_ext}'. "
                f"Accepted: {accepted}."
            ),
        )

    base = re.sub(r'\.[^.]+$', '', src_name) or 'paper_spec'
    target_name = f'{base}.pdf'
    warnings: List[str] = []

    try:
        # Read bytes once (Django chunks() if streaming).
        django_file.seek(0)
        raw = django_file.read()
        django_file.seek(0)

        engine = desc['engine']

        if engine == 'passthrough':
            return NormalizedFile(
                success=True,
                file=ContentFile(raw, name=src_name),
                filename=src_name,
                mime='application/pdf',
                original_extension=src_ext,
                original_mime=src_mime,
                conversion_engine='passthrough',
                conversion_notes='Input was already a PDF.',
            )

        if engine == 'image_to_pdf':
            pdf_bytes = _image_to_pdf(raw, src_ext)
            notes = f'Wrapped {src_ext.upper()} image into a single-page PDF.'
        elif engine == 'docx_to_pdf':
            pdf_bytes = _docx_to_pdf(raw)
            notes = 'Extracted DOCX text + tables and rendered to PDF.'
            warnings.append('Inline images inside Word documents are not preserved — only text is extracted.')
        elif engine == 'xlsx_to_pdf':
            pdf_bytes = _xlsx_to_pdf(raw)
            notes = 'Flattened Excel sheets to pipe-separated text PDF.'
            warnings.append(
                f"Rows capped at {NORMALIZER_CONFIG['xlsx_max_rows_per_sheet']} per sheet "
                f"and columns capped at {NORMALIZER_CONFIG['xlsx_max_cols']}."
            )
        elif engine == 'text_to_pdf':
            try:
                text = raw.decode('utf-8', errors='replace')
            except Exception:
                text = raw.decode('latin-1', errors='replace')
            pdf_bytes = _text_to_pdf(text, title=base)
            notes = 'Rendered plain text into PDF.'
        else:
            return NormalizedFile(
                success=False,
                original_extension=src_ext,
                original_mime=src_mime,
                error_message=f"Internal: no handler for engine '{engine}'.",
            )

        return NormalizedFile(
            success=True,
            file=ContentFile(pdf_bytes, name=target_name),
            filename=target_name,
            mime='application/pdf',
            original_extension=src_ext,
            original_mime=src_mime,
            conversion_engine=engine,
            conversion_notes=notes,
            warnings=warnings,
        )

    except Exception as exc:
        logger.exception('[SpecNormalize] Conversion failed for %s', src_name)
        return NormalizedFile(
            success=False,
            original_extension=src_ext,
            original_mime=src_mime,
            conversion_engine=desc.get('engine', ''),
            error_message=f'Could not convert .{src_ext} to PDF: {exc}',
        )

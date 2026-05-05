"""
Multi-format document text extractor (soft-coded).

Centralised utility used by every electrical datasheet generator
(transformer / DG set / switchgear / …) so the **same set of accepted
formats** is enforced in views, generators and the frontend.

Supported formats are declared in `SUPPORTED_FORMATS` below.  To add a new
format:

  1. Append the extension to `SUPPORTED_FORMATS`.
  2. Implement an `_extract_<ext>` handler.
  3. Wire it into `_HANDLERS`.

Frontend `DEFAULT_ACCEPTED_FORMATS` must mirror `SUPPORTED_FORMATS`.
"""
from __future__ import annotations

import csv
import io
import logging
from typing import Callable, Dict, Iterable, Tuple

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Soft-coded format registry
# ─────────────────────────────────────────────────────────────────────────────
SUPPORTED_FORMATS: Tuple[str, ...] = (
    ".pdf",
    ".xlsx", ".xlsm", ".xls",
    ".docx",
    ".csv",
    ".txt",
    ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif",
)

SUPPORTED_FORMATS_LABEL = "PDF · Excel · Word · CSV · TXT · Images"

MIN_TEXT_LEN = 20  # below this, treat extraction as failed


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────
def is_supported(filename: str) -> bool:
    if not filename:
        return False
    name = filename.lower()
    return any(name.endswith(ext) for ext in SUPPORTED_FORMATS)


def get_extension(filename: str) -> str:
    name = (filename or "").lower()
    for ext in SUPPORTED_FORMATS:
        if name.endswith(ext):
            return ext
    return ""


def extract_text(uploaded_file) -> str:
    """
    Extract plain text from any supported uploaded file.

    Falls back gracefully — if a specific parser is unavailable or fails on
    a particular document, returns whatever text could be recovered (may be
    empty).  Callers should check `len(text) >= MIN_TEXT_LEN`.
    """
    name = getattr(uploaded_file, "name", "") or ""
    ext = get_extension(name)
    handler = _HANDLERS.get(ext)
    if handler is None:
        logger.warning("[DocumentExtractor] Unsupported extension for %s", name)
        return ""
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    try:
        text = handler(uploaded_file) or ""
        logger.info("[DocumentExtractor] %s → %d chars (%s)", name, len(text), ext)
        return text
    except Exception as exc:
        logger.error("[DocumentExtractor] %s extraction failed: %s", ext, exc, exc_info=True)
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Per-format handlers
# ─────────────────────────────────────────────────────────────────────────────
def _extract_pdf(file_obj) -> str:
    import PyPDF2
    file_obj.seek(0)
    reader = PyPDF2.PdfReader(file_obj)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n".join(parts).strip()
    # If text-based extraction produced nothing useful, try OCR on rasterised pages.
    if len(text) < MIN_TEXT_LEN:
        ocr_text = _ocr_pdf(file_obj)
        if ocr_text:
            text = ocr_text
    return text


def _ocr_pdf(file_obj) -> str:
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except Exception:
        return ""
    try:
        file_obj.seek(0)
        images = convert_from_bytes(file_obj.read())
        return "\n".join(pytesseract.image_to_string(img) for img in images).strip()
    except Exception as exc:
        logger.warning("[DocumentExtractor] PDF OCR fallback failed: %s", exc)
        return ""


def _extract_xlsx(file_obj) -> str:
    import openpyxl
    file_obj.seek(0)
    data = io.BytesIO(file_obj.read())
    wb = openpyxl.load_workbook(data, data_only=True, read_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                lines.append("\t".join(cells))
    return "\n".join(lines).strip()


def _extract_xls(file_obj) -> str:
    # Legacy .xls — try pandas (xlrd backend) first
    try:
        import pandas as pd
        file_obj.seek(0)
        sheets = pd.read_excel(file_obj, sheet_name=None, dtype=str)
        out = []
        for name, df in sheets.items():
            out.append(f"# Sheet: {name}")
            out.append(df.fillna("").to_csv(sep="\t", index=False))
        return "\n".join(out).strip()
    except Exception as exc:
        logger.warning("[DocumentExtractor] .xls extraction failed: %s", exc)
        return ""


def _extract_docx(file_obj) -> str:
    from docx import Document
    file_obj.seek(0)
    doc = Document(io.BytesIO(file_obj.read()))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_csv(file_obj) -> str:
    file_obj.seek(0)
    raw = file_obj.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(raw))
    return "\n".join("\t".join(row) for row in reader).strip()


def _extract_txt(file_obj) -> str:
    file_obj.seek(0)
    raw = file_obj.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")
    return raw.strip()


def _extract_image(file_obj) -> str:
    try:
        from PIL import Image
        import pytesseract
    except Exception:
        logger.warning("[DocumentExtractor] OCR libraries not installed for image extraction")
        return ""
    try:
        file_obj.seek(0)
        img = Image.open(io.BytesIO(file_obj.read()))
        return pytesseract.image_to_string(img).strip()
    except Exception as exc:
        logger.warning("[DocumentExtractor] Image OCR failed: %s", exc)
        return ""


_HANDLERS: Dict[str, Callable] = {
    ".pdf":  _extract_pdf,
    ".xlsx": _extract_xlsx,
    ".xlsm": _extract_xlsx,
    ".xls":  _extract_xls,
    ".docx": _extract_docx,
    ".csv":  _extract_csv,
    ".txt":  _extract_txt,
    ".png":  _extract_image,
    ".jpg":  _extract_image,
    ".jpeg": _extract_image,
    ".bmp":  _extract_image,
    ".tiff": _extract_image,
    ".tif":  _extract_image,
}

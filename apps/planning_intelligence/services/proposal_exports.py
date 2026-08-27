"""PDF and Word exports for frozen enterprise technical proposals."""
from __future__ import annotations

import io
import re
from xml.sax.saxutils import escape


def _safe_filename(proposal, extension):
    base = re.sub(r'[^A-Za-z0-9_-]+', '-', proposal.proposal_number).strip('-') or 'proposal'
    return f'{base}-Rev-{proposal.revision}.{extension}'


def _rows(section):
    rows = section.get('data') or []
    if not rows or not isinstance(rows[0], dict):
        return [], []
    headers = list(rows[0].keys())[:7]
    return headers, rows


def _proposal_pdf_legacy(proposal):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    stream = io.BytesIO()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ProposalTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=23, leading=28, textColor=colors.HexColor('#273B5A'), alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name='ProposalH1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=15, leading=18, textColor=colors.HexColor('#123FD1'), spaceBefore=3, spaceAfter=10))
    styles.add(ParagraphStyle(name='ProposalBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14.5, textColor=colors.HexColor('#111827'), spaceAfter=7))
    styles.add(ParagraphStyle(name='ProposalTable', parent=styles['BodyText'], fontName='Helvetica', fontSize=7.5, leading=9.5, textColor=colors.HexColor('#1F2937')))

    def header_footer(canvas, document):
        canvas.saveState()
        classification = proposal.branding.get('confidentiality', 'Confidential').title()
        if document.page > 1:
            canvas.setFillColor(colors.HexColor('#111827'))
            canvas.setFont('Helvetica-Bold', 7.5)
            canvas.drawString(16 * mm, 287 * mm, 'Technical Proposal')
            canvas.setFont('Helvetica', 7.5)
            canvas.drawString(50 * mm, 287 * mm, proposal.proposal_number)
            canvas.drawString(16 * mm, 282.5 * mm, classification)
            canvas.drawString(50 * mm, 282.5 * mm, f'Rev {proposal.revision} / {proposal.get_status_display()}')
            canvas.setFont('Helvetica-Bold', 13)
            canvas.setFillColor(colors.HexColor('#273B5A'))
            canvas.drawRightString(194 * mm, 285 * mm, 'REJLERS')
            canvas.setStrokeColor(colors.HexColor('#CBD5E1'))
            canvas.line(16 * mm, 278 * mm, 194 * mm, 278 * mm)
        canvas.setStrokeColor(colors.HexColor('#8793A5'))
        canvas.line(16 * mm, 16 * mm, 194 * mm, 16 * mm)
        canvas.setFillColor(colors.HexColor('#65748A'))
        canvas.setFont('Helvetica-Bold', 5.8)
        canvas.drawString(16 * mm, 12.8 * mm, 'Rejlers International Engineering Solutions AB')
        canvas.setFont('Helvetica', 5.6)
        canvas.drawString(16 * mm, 10.2 * mm, 'Millennium Tower, 13th Floor, Hamdan Street, P.O. Box 39317, Abu Dhabi, United Arab Emirates')
        canvas.drawString(16 * mm, 7.7 * mm, 'Tel: +971 2 639 7449  |  www.rejlers.ae')
        canvas.drawRightString(194 * mm, 10.2 * mm, f'{proposal.proposal_number} · Rev {proposal.revision}')
        canvas.drawRightString(194 * mm, 7.7 * mm, f'Page {document.page}')
        canvas.restoreState()

    doc = SimpleDocTemplate(stream, pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm, topMargin=26 * mm, bottomMargin=22 * mm, title=proposal.title)
    story = [Spacer(1, 42 * mm), Paragraph(escape(proposal.title), styles['ProposalTitle']), Spacer(1, 8 * mm), Paragraph(escape(proposal.project.name), styles['Title']), Spacer(1, 5 * mm), Paragraph(escape(f'Prepared for {proposal.client_name or proposal.project.client or "Client"}'), styles['Heading2']), Spacer(1, 25 * mm)]
    control = [
        [Paragraph('<b>Proposal Number</b>', styles['BodyText']), proposal.proposal_number],
        [Paragraph('<b>Revision / Status</b>', styles['BodyText']), f'{proposal.revision} / {proposal.get_status_display()}'],
        [Paragraph('<b>Tender / RFT Reference</b>', styles['BodyText']), proposal.opportunity_reference or 'Not specified'],
        [Paragraph('<b>Client Reference</b>', styles['BodyText']), proposal.client_reference or 'Not specified'],
        [Paragraph('<b>Submission Date</b>', styles['BodyText']), str(proposal.submission_date or 'Not specified')],
        [Paragraph('<b>Offer Validity</b>', styles['BodyText']), f'{proposal.validity_days} days'],
        [Paragraph('<b>Schedule Version</b>', styles['BodyText']), str(proposal.schedule_version.version)],
    ]
    table = Table(control, colWidths=[45 * mm, 90 * mm])
    table.setStyle(TableStyle([('GRID', (0, 0), (-1, -1), .5, colors.HexColor('#CBD5E1')), ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#EEF2FF')), ('VALIGN', (0, 0), (-1, -1), 'TOP'), ('PADDING', (0, 0), (-1, -1), 6)]))
    story += [table, PageBreak(), Paragraph('Table of Contents', styles['ProposalH1'])]
    current_group = None
    for item in proposal.sections:
        if not item.get('included', True) or item.get('key') in ('cover', 'contents'):
            continue
        group = item.get('group') or 'Technical Proposal'
        if group != current_group:
            story += [Spacer(1, 3 * mm), Paragraph(f'<b>{escape(group)}</b>', styles['BodyText'])]
            current_group = group
        label = f'{item.get("number", "")} {item.get("title", "Section")}'.strip()
        story.append(Paragraph(escape(label), styles['ProposalBody']))
    story.append(PageBreak())

    current_group = None
    for section in proposal.sections:
        if not section.get('included', True) or section.get('key') in ('cover', 'contents'):
            continue
        group = section.get('group') or 'Technical Proposal'
        if group != current_group:
            story += [Paragraph(escape(group.upper()), styles['Heading3'])]
            current_group = group
        heading = f'{section.get("number", "")} {section.get("title", "Section")}'.strip()
        story += [Paragraph(escape(heading), styles['ProposalH1'])]
        for paragraph in str(section.get('content') or '').split('\n'):
            if paragraph.strip():
                story.append(Paragraph(escape(paragraph), styles['ProposalBody']))
        headers, rows = _rows(section)
        if headers:
            data = [[Paragraph(f'<b>{escape(str(header).replace("_", " ").title())}</b>', styles['ProposalTable']) for header in headers]]
            for row in rows:
                data.append([Paragraph(escape(str(row.get(header, '') or '')), styles['ProposalTable']) for header in headers])
            grid = Table(data, repeatRows=1, hAlign='LEFT')
            grid.setStyle(TableStyle([('GRID', (0, 0), (-1, -1), .35, colors.HexColor('#CBD5E1')), ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')), ('VALIGN', (0, 0), (-1, -1), 'TOP'), ('FONTSIZE', (0, 0), (-1, -1), 7), ('PADDING', (0, 0), (-1, -1), 3)]))
            story += [Spacer(1, 4 * mm), grid]
        story.append(PageBreak())
    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
    return stream.getvalue(), 'application/pdf', _safe_filename(proposal, 'pdf')


def proposal_pdf(proposal):
    """Generate the controlled, fully paginated A4 proposal with a real TOC."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageBreak, PageTemplate, Paragraph, Spacer,
        Table, TableStyle,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    stream = io.BytesIO()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ProposalTitle', parent=styles['Title'], fontName='Helvetica-Bold',
        fontSize=23, leading=28, textColor=colors.HexColor('#273B5A'),
        alignment=TA_CENTER, spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name='ProposalH1', parent=styles['Heading1'], fontName='Helvetica-Bold',
        fontSize=15, leading=18, textColor=colors.HexColor('#123FD1'),
        spaceBefore=3, spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name='ProposalBody', parent=styles['BodyText'], fontName='Helvetica',
        fontSize=10, leading=14.5, textColor=colors.HexColor('#111827'), spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name='ProposalTable', parent=styles['BodyText'], fontName='Helvetica',
        fontSize=7.5, leading=9.5, textColor=colors.HexColor('#1F2937'),
    ))
    styles.add(ParagraphStyle(
        name='ProposalTOC', parent=styles['BodyText'], fontName='Helvetica',
        fontSize=9.5, leading=14, leftIndent=0, firstLineIndent=0,
        textColor=colors.HexColor('#111827'), alignment=TA_LEFT,
    ))

    class ProposalDocTemplate(BaseDocTemplate):
        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph) and flowable.style.name == 'ProposalH1':
                title = flowable.getPlainText()
                if title == 'Table Content':
                    return
                key = f'proposal-section-{self.seq.nextf("section")}'
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(title, key, level=0, closed=False)
                self.notify('TOCEntry', (0, title, self.page, key))

    def page_chrome(canvas, document):
        canvas.saveState()
        classification = proposal.branding.get('confidentiality', 'Confidential').title()
        brand_blue = colors.HexColor('#0870AA')
        brand_text = colors.HexColor('#3275B6')

        # The border, header, and footer deliberately mirror the new PO form shell.
        canvas.setStrokeColor(colors.HexColor('#94A3B8'))
        canvas.setLineWidth(.65)
        canvas.rect(8 * mm, 6 * mm, 194 * mm, 285 * mm, stroke=1, fill=0)

        canvas.setFillColor(brand_text)
        canvas.setFont('Helvetica-Bold', 8)
        canvas.drawString(14 * mm, 284.5 * mm, 'TECHNICAL PROPOSAL')
        canvas.setFont('Helvetica-Bold', 7)
        canvas.drawString(14 * mm, 280.3 * mm, proposal.proposal_number)
        canvas.setFont('Helvetica', 6.2)
        canvas.drawString(14 * mm, 276.5 * mm, f'{classification}  |  Rev {proposal.revision} / {proposal.get_status_display()}')
        canvas.setFont('Helvetica-Bold', 14)
        canvas.drawRightString(196 * mm, 284 * mm, 'REJLERS')
        canvas.setFont('Helvetica-Bold', 5.5)
        canvas.drawRightString(196 * mm, 279.7 * mm, 'HOME OF THE LEARNING MINDS')
        canvas.setStrokeColor(colors.HexColor('#CBD5E1'))
        canvas.line(14 * mm, 273 * mm, 196 * mm, 273 * mm)

        canvas.setFillColor(brand_blue)
        canvas.rect(8 * mm, 15 * mm, 194 * mm, 7 * mm, stroke=0, fill=1)
        canvas.setFillColor(colors.white)
        canvas.setFont('Helvetica-Bold', 5.2)
        canvas.drawString(14 * mm, 17.7 * mm, 'REJLERS')
        canvas.drawCentredString(105 * mm, 18.1 * mm, 'HOME OF THE LEARNING MINDS')
        canvas.drawRightString(196 * mm, 17.7 * mm, 'RADAI')
        canvas.setFillColor(brand_text)
        canvas.setFont('Helvetica-Bold', 5.2)
        canvas.drawString(14 * mm, 11.8 * mm, 'Rejlers International Engineering Solutions')
        canvas.setFont('Helvetica', 4.9)
        canvas.drawString(14 * mm, 9.2 * mm, 'Millennium Tower, 13th Floor, Hamdan Street, P.O. Box 39317, Abu Dhabi, UAE')
        canvas.drawString(14 * mm, 7.1 * mm, 'Tel: +971 2 639 7449  |  www.rejlers.ae')
        canvas.drawRightString(196 * mm, 10.2 * mm, f'{proposal.proposal_number}  |  Rev {proposal.revision}')
        canvas.drawRightString(196 * mm, 7.4 * mm, f'Page {document.page}')
        canvas.restoreState()

    doc = ProposalDocTemplate(stream, pagesize=A4, title=proposal.title)
    frame = Frame(
        14 * mm, 23 * mm, 182 * mm, 247 * mm,
        leftPadding=2 * mm, rightPadding=2 * mm,
        topPadding=2 * mm, bottomPadding=2 * mm, id='proposal-frame',
    )
    doc.addPageTemplates([PageTemplate(id='proposal', frames=[frame], onPage=page_chrome)])

    story = [
        Spacer(1, 36 * mm), Paragraph(escape(proposal.title), styles['ProposalTitle']),
        Spacer(1, 8 * mm), Paragraph(escape(proposal.project.name), styles['Title']),
        Spacer(1, 5 * mm),
        Paragraph(escape(f'Prepared for {proposal.client_name or proposal.project.client or "Client"}'), styles['Heading2']),
        Spacer(1, 22 * mm),
    ]
    control = [
        [Paragraph('<b>Proposal Number</b>', styles['BodyText']), proposal.proposal_number],
        [Paragraph('<b>Revision / Status</b>', styles['BodyText']), f'{proposal.revision} / {proposal.get_status_display()}'],
        [Paragraph('<b>Tender / RFT Reference</b>', styles['BodyText']), proposal.opportunity_reference or 'Not specified'],
        [Paragraph('<b>Client Reference</b>', styles['BodyText']), proposal.client_reference or 'Not specified'],
        [Paragraph('<b>Submission Date</b>', styles['BodyText']), str(proposal.submission_date or 'Not specified')],
        [Paragraph('<b>Offer Validity</b>', styles['BodyText']), f'{proposal.validity_days} days'],
        [Paragraph('<b>Schedule Version</b>', styles['BodyText']), str(proposal.schedule_version.version)],
    ]
    control_table = Table(control, colWidths=[45 * mm, 90 * mm])
    control_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), .5, colors.HexColor('#CBD5E1')),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#EEF2FF')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'), ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story += [control_table, PageBreak(), Paragraph('Table Content', styles['ProposalH1'])]

    contents = TableOfContents()
    contents.levelStyles = [styles['ProposalTOC']]
    contents.dotsMinLevel = 0
    toc_header = Table([
        [Paragraph('<b>Table Content</b>', styles['ProposalBody']), Paragraph('<b>Page</b>', styles['ProposalBody'])],
    ], colWidths=[160 * mm, 16 * mm])
    toc_header.setStyle(TableStyle([
        ('LINEBELOW', (0, 0), (-1, -1), .6, colors.HexColor('#94A3B8')),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
    ]))
    story += [toc_header, Spacer(1, 3 * mm), contents, PageBreak()]

    printable_sections = [
        section for section in proposal.sections
        if section.get('included', True) and section.get('key') not in ('cover', 'contents')
    ]
    current_group = None
    for section_index, section in enumerate(printable_sections):
        group = section.get('group') or 'Technical Proposal'
        if group != current_group:
            story += [Paragraph(escape(group.upper()), styles['Heading3'])]
            current_group = group
        heading = f'{section.get("number", "")} {section.get("title", "Section")}'.strip()
        story += [Paragraph(escape(heading), styles['ProposalH1'])]
        for paragraph in str(section.get('content') or '').split('\n'):
            if paragraph.strip():
                story.append(Paragraph(escape(paragraph), styles['ProposalBody']))
        headers, rows = _rows(section)
        if headers:
            data = [[Paragraph(f'<b>{escape(str(header).replace("_", " ").title())}</b>', styles['ProposalTable']) for header in headers]]
            for row in rows:
                data.append([Paragraph(escape(str(row.get(header, '') or '')), styles['ProposalTable']) for header in headers])
            grid = Table(data, repeatRows=1, hAlign='LEFT')
            grid.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), .35, colors.HexColor('#CBD5E1')),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), 7), ('PADDING', (0, 0), (-1, -1), 3),
            ]))
            story += [Spacer(1, 4 * mm), grid]
        if section_index < len(printable_sections) - 1:
            story.append(PageBreak())

    doc.multiBuild(story)
    return stream.getvalue(), 'application/pdf', _safe_filename(proposal, 'pdf')


def proposal_docx(proposal):
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(.82); section.bottom_margin = Inches(.72)
    section.left_margin = Inches(.63); section.right_margin = Inches(.63)
    section.header_distance = Inches(.25); section.footer_distance = Inches(.2)
    section.different_first_page_header_footer = False

    # Match the controlled A4 border used by the browser and PDF renderers.
    section_properties = section._sectPr
    page_borders = OxmlElement('w:pgBorders')
    page_borders.set(qn('w:offsetFrom'), 'page')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '18')
        border.set(qn('w:color'), '94A3B8')
        page_borders.append(border)
    section_properties.append(page_borders)

    normal = document.styles['Normal']
    normal.font.name = 'Arial'; normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in [
        ('Title', 23, '273B5A'), ('Heading 1', 15, '123FD1'),
        ('Heading 2', 12, '273B5A'), ('Heading 3', 10, '536783'),
    ]:
        style = document.styles[style_name]
        style.font.name = 'Arial'; style.font.size = Pt(size); style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header_table = section.header.add_table(rows=2, cols=3, width=Inches(7.0))
    header_values = [
        ('TECHNICAL PROPOSAL', proposal.proposal_number, 'REJLERS'),
        (proposal.branding.get('confidentiality', 'Confidential').title(), f'Rev {proposal.revision} / {proposal.get_status_display()}', 'HOME OF THE LEARNING MINDS'),
    ]
    for row_index, values in enumerate(header_values):
        for cell_index, value in enumerate(values):
            cell = header_table.rows[row_index].cells[cell_index]
            cell.text = str(value)
            paragraph = cell.paragraphs[0]
            if cell_index == 2: paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.name = 'Arial'; run.font.size = Pt(8 if cell_index < 2 else (13 if row_index == 0 else 5.5))
                run.font.bold = row_index == 0 and cell_index in (0, 2)
                run.font.color.rgb = RGBColor.from_string('3275B6')

    def add_page_field(paragraph):
        paragraph.add_run('Page ')
        run = paragraph.add_run()
        begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
        instruction = OxmlElement('w:instrText'); instruction.set(qn('xml:space'), 'preserve'); instruction.text = ' PAGE '
        end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
        run._r.extend([begin, instruction, end])

    def populate_footer(footer):
        brand_bar = footer.add_table(rows=1, cols=3, width=Inches(7.0))
        for index, value in enumerate(('REJLERS', 'HOME OF THE LEARNING MINDS', 'RADAI')):
            cell = brand_bar.rows[0].cells[index]
            cell.text = value
            shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), '0870AA')
            cell._tc.get_or_add_tcPr().append(shading)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Arial'; run.font.size = Pt(5.5); run.font.bold = True
                run.font.color.rgb = RGBColor.from_string('FFFFFF')
        table = footer.add_table(rows=1, cols=2, width=Inches(7.0))
        left, right = table.rows[0].cells
        left.text = 'Rejlers International Engineering Solutions AB\nMillennium Tower, 13th Floor, Hamdan Street, P.O. Box 39317, Abu Dhabi, UAE\nTel: +971 2 639 7449  |  www.rejlers.ae'
        right.text = f'{proposal.proposal_number} · Rev {proposal.revision}\n'
        add_page_field(right.paragraphs[-1])
        right.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in (left, right):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'; run.font.size = Pt(6.5)
                    run.font.color.rgb = RGBColor.from_string('65748A')

    populate_footer(section.footer)
    title = document.add_heading(proposal.title, 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(proposal.project.name); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client = document.add_paragraph(f'Prepared for {proposal.client_name or proposal.project.client or "Client"}'); client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    control = document.add_table(rows=0, cols=2); control.style = 'Table Grid'
    for label, value in [('Proposal Number', proposal.proposal_number), ('Revision', proposal.revision), ('Status', proposal.get_status_display()), ('Tender / RFT Reference', proposal.opportunity_reference or 'Not specified'), ('Client Reference', proposal.client_reference or 'Not specified'), ('Submission Date', proposal.submission_date or 'Not specified'), ('Offer Validity', f'{proposal.validity_days} days'), ('Schedule Version', proposal.schedule_version.version)]:
        cells = control.add_row().cells; cells[0].text = str(label); cells[1].text = str(value)
    document.add_page_break()
    document.add_heading('Table Content', level=1)
    toc_header = document.add_table(rows=1, cols=2)
    toc_cells = toc_header.rows[0].cells
    toc_cells[0].text = 'Table Content'
    toc_cells[1].text = 'Page'
    toc_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in toc_cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True; run.font.name = 'Arial'; run.font.size = Pt(9)
    toc_paragraph = document.add_paragraph()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin'); begin.set(qn('w:dirty'), 'true')
    instruction = OxmlElement('w:instrText'); instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t'); placeholder.text = 'Right-click and select Update Field if page numbers are not displayed.'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    toc_run = toc_paragraph.add_run()._r
    toc_run.extend([begin, instruction, separate, placeholder, end])
    document.add_page_break()
    current_group = None
    for item in proposal.sections:
        if not item.get('included', True) or item.get('key') in ('cover', 'contents'):
            continue
        group = item.get('group') or 'Technical Proposal'
        if group != current_group:
            document.add_heading(group, level=1); current_group = group
        heading = f'{item.get("number", "")} {item.get("title", "Section")}'.strip()
        document.add_heading(heading, level=2)
        for paragraph in str(item.get('content') or '').split('\n'):
            if paragraph.strip(): document.add_paragraph(paragraph)
        headers, rows = _rows(item)
        if headers:
            table = document.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
            for index, header in enumerate(headers): table.rows[0].cells[index].text = str(header).replace('_', ' ').title()
            for row in rows:
                cells = table.add_row().cells
                for index, header in enumerate(headers): cells[index].text = str(row.get(header, '') or '')
        document.add_page_break()
    for section in document.sections:
        for table in section.header.tables + section.footer.tables + section.first_page_footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'; run.font.size = Pt(8)
    settings = document.settings._element
    update_fields = OxmlElement('w:updateFields'); update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)
    stream = io.BytesIO(); document.save(stream)
    return stream.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', _safe_filename(proposal, 'docx')


def generate_proposal_export(proposal, export_format):
    if export_format == 'pdf': return proposal_pdf(proposal)
    if export_format == 'docx': return proposal_docx(proposal)
    raise ValueError('Unsupported proposal export format.')

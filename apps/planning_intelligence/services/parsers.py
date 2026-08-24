"""
Document parsing service — best-effort text extraction (MODULE 2).

Modular parser interface: each file category / extension routes to a small
extractor function. Every extractor is wrapped so a parsing failure never
crashes the request — it degrades to an empty string + low confidence score,
and the caller (views.py) records parse_status='failed' with the error.
"""
from __future__ import annotations

import csv
import io
import logging

logger = logging.getLogger(__name__)

# Soft-coded per-extension max characters kept (keeps JSONField/TextField sane).
MAX_EXTRACTED_CHARS = 400_000
MAX_OCR_PAGES = 50
MIN_PDF_TEXT_CHARS = 80

# Binary/image extensions have no extractable text — never route them through
# _extract_plain_text, which would decode raw bytes as UTF-8 and can produce
# NUL (0x00) characters that Postgres TEXT columns reject outright.
_BINARY_EXTENSIONS = {
    'gif', 'bmp', 'webp', 'ico', 'svg', 'zip', 'rar', '7z', 'exe', 'dll',
}


def _truncate(text: str) -> str:
    # Defensive: strip NUL bytes from any extractor's output — Postgres TEXT
    # columns reject them and would otherwise crash the save() call.
    text = text.replace('\x00', '')
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS] + '\n...[truncated]'
    return text


def _extract_pdf(file_obj) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or '')
        joined = '\f'.join(text_parts).strip()
        if len(joined) >= MIN_PDF_TEXT_CHARS:
            return joined
    except Exception as exc:  # noqa: BLE001
        logger.info('pdfplumber failed (%s); falling back to PyPDF2', exc)

    try:
        file_obj.seek(0)
        from PyPDF2 import PdfReader
        reader = PdfReader(file_obj)
        joined = '\f'.join((page.extract_text() or '') for page in reader.pages).strip()
        if len(joined) >= MIN_PDF_TEXT_CHARS:
            return joined
    except Exception as exc:  # noqa: BLE001
        logger.warning('PyPDF2 fallback also failed: %s', exc)
    return _extract_pdf_ocr(file_obj)


def _extract_pdf_ocr(file_obj) -> str:
    """Bounded optional OCR fallback for image-only/scanned PDF references."""
    try:
        import fitz
        import pytesseract
        from PIL import Image

        file_obj.seek(0)
        document = fitz.open(stream=file_obj.read(), filetype='pdf')
        pages = []
        for index, page in enumerate(document):
            if index >= MAX_OCR_PAGES:
                break
            pixmap = page.get_pixmap(matrix=fitz.Matrix(200 / 72, 200 / 72), alpha=False)
            image = Image.open(io.BytesIO(pixmap.tobytes('png')))
            text = pytesseract.image_to_string(image, lang='eng', config='--oem 3 --psm 6').strip()
            pages.append(f'--- OCR Page: {index + 1} ---\n{text}')
        return '\f'.join(pages).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning('Scanned PDF OCR fallback unavailable or failed: %s', exc)
        return ''


def _extract_xlsx(file_obj) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_obj, data_only=True, read_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f'--- Sheet: {ws.title} ---')
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    lines.append(' | '.join(cells))
        return '\n'.join(lines)
    except Exception as exc:  # noqa: BLE001
        logger.warning('openpyxl extraction failed: %s', exc)
        return ''


def _extract_csv(file_obj) -> str:
    try:
        raw = file_obj.read()
        text = raw.decode('utf-8', errors='ignore') if isinstance(raw, bytes) else raw
        reader = csv.reader(io.StringIO(text))
        return '\n'.join(' | '.join(row) for row in reader)
    except Exception as exc:  # noqa: BLE001
        logger.warning('csv extraction failed: %s', exc)
        return ''


def _extract_docx(file_obj) -> str:
    try:
        import docx
        document = docx.Document(file_obj)
        lines = [p.text for p in document.paragraphs if p.text.strip()]
        for table_index, table in enumerate(document.tables, start=1):
            lines.append(f'--- Table: {table_index} ---')
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(' | '.join(cells))
        return '\n'.join(lines)
    except Exception as exc:  # noqa: BLE001
        logger.warning('docx extraction failed: %s', exc)
        return ''


def _extract_plain_text(file_obj) -> str:
    try:
        raw = file_obj.read()
        return raw.decode('utf-8', errors='ignore') if isinstance(raw, bytes) else str(raw)
    except Exception as exc:  # noqa: BLE001
        logger.warning('plain text extraction failed: %s', exc)
        return ''


def _extract_image_ocr(file_obj) -> str:
    try:
        import pytesseract
        from PIL import Image
        file_obj.seek(0)
        text = pytesseract.image_to_string(Image.open(file_obj), lang='eng', config='--oem 3 --psm 6').strip()
        return f'--- OCR Page: 1 ---\n{text}' if text else ''
    except Exception as exc:  # noqa: BLE001
        logger.warning('Image OCR extraction unavailable or failed: %s', exc)
        return ''


_EXTRACTORS = {
    'pdf': _extract_pdf,
    'xlsx': _extract_xlsx,
    'xlsm': _extract_xlsx,
    'xls': _extract_xlsx,
    'csv': _extract_csv,
    'docx': _extract_docx,
    'txt': _extract_plain_text,
    'xer': _extract_plain_text,  # Primavera native export is plain text (tab-delimited)
    'png': _extract_image_ocr,
    'jpg': _extract_image_ocr,
    'jpeg': _extract_image_ocr,
    'tif': _extract_image_ocr,
    'tiff': _extract_image_ocr,
}


def extract_text(file_field, original_filename: str) -> tuple[str, float]:
    """
    Returns (extracted_text, confidence_score). confidence_score is a coarse
    heuristic (0.0-1.0) based on extracted content length — NOT a real ML
    confidence — kept intentionally simple per MVP scope.
    """
    ext = (original_filename.rsplit('.', 1)[-1] if '.' in original_filename else '').lower()

    if ext in _BINARY_EXTENSIONS:
        return '', 0.0

    extractor = _EXTRACTORS.get(ext, _extract_plain_text)

    try:
        file_field.seek(0)
    except Exception:  # noqa: BLE001
        pass

    text = extractor(file_field) or ''
    text = _truncate(text.strip())

    if not text:
        return '', 0.0
    confidence = min(1.0, 0.3 + len(text) / 20000)
    return text, round(confidence, 2)
